import os
import tempfile
import pytesseract
import fitz  # PyMuPDF
from PIL import Image
from pdf2image import convert_from_path
import google.generativeai as genai
import logging
import json
import re
from typing import Dict, List, Optional
import docx
from docx import Document
import argparse
import sys
from pathlib import Path
import glob

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


# -------------------------------
# 1. CONFIGURE GEMINI API
# -------------------------------
def configure_gemini(api_key: str = None):
    """Configure Gemini API with provided key."""
    if not api_key:
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            raise ValueError(
                "Please provide Gemini API key either as parameter or set GEMINI_API_KEY environment variable")

    genai.configure(api_key=api_key)
    logging.info("Gemini API configured successfully")


# -------------------------------
# 2. FILE VALIDATION AND DISCOVERY
# -------------------------------
def get_supported_files(input_path: str) -> List[str]:
    """Get list of supported files from input path (file or directory)."""
    supported_extensions = {'.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.doc', '.docx', '.txt'}

    if os.path.isfile(input_path):
        # Single file
        if Path(input_path).suffix.lower() in supported_extensions:
            return [input_path]
        else:
            raise ValueError(f"Unsupported file format: {Path(input_path).suffix}")

    elif os.path.isdir(input_path):
        # Directory - find all supported files
        files = []
        for ext in supported_extensions:
            pattern = os.path.join(input_path, f"*{ext}")
            files.extend(glob.glob(pattern))
            # Also check uppercase extensions
            pattern = os.path.join(input_path, f"*{ext.upper()}")
            files.extend(glob.glob(pattern))

        if not files:
            raise ValueError(f"No supported files found in directory: {input_path}")

        return sorted(files)

    else:
        # Check if it's a glob pattern
        files = glob.glob(input_path)
        if files:
            supported_files = []
            for file in files:
                if Path(file).suffix.lower() in supported_extensions:
                    supported_files.append(file)

            if not supported_files:
                raise ValueError(f"No supported files found matching pattern: {input_path}")

            return sorted(supported_files)
        else:
            raise FileNotFoundError(f"File or directory not found: {input_path}")


# -------------------------------
# 3. ENHANCED TEXT EXTRACTION
# -------------------------------
def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats with improved error handling."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    text = ""
    file_ext = Path(file_path).suffix.lower()

    try:
        if file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
            text = extract_text_from_image(file_path)
        elif file_ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            text = extract_text_from_doc(file_path)
        elif file_ext == '.txt':
            text = extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        logging.info(f"Successfully extracted {len(text)} characters from {file_path}")
        return text.strip()

    except Exception as e:
        logging.error(f"Text extraction failed for {file_path}: {e}")
        raise


def extract_text_from_image(file_path: str) -> str:
    """Extract text from image files using OCR."""
    try:
        img = Image.open(file_path)
        if img.mode != 'RGB':
            img = img.convert('RGB')

        # Try different OCR configurations for better results
        ocr_configs = [
            '--psm 6',  # Uniform block of text
            '--psm 4',  # Single column of text
            '--psm 3',  # Fully automatic page segmentation
        ]

        text = ""
        for config in ocr_configs:
            try:
                text = pytesseract.image_to_string(img, config=config)
                if text.strip():  # If we got text, use it
                    break
            except:
                continue

        return text
    except Exception as e:
        logging.error(f"Image OCR failed: {e}")
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files with OCR fallback."""
    text = ""

    try:
        # First try direct text extraction
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                text += page_text + "\n"

        # If direct extraction didn't yield much text, try OCR
        if not text.strip() or len(text.strip()) < 100:
            logging.info("Direct PDF text extraction yielded little text, trying OCR...")
            try:
                with tempfile.TemporaryDirectory() as temp_dir:
                    images = convert_from_path(file_path, output_folder=temp_dir, dpi=300)
                    ocr_text = ""
                    for i, img in enumerate(images):
                        page_text = pytesseract.image_to_string(img, config='--psm 6')
                        ocr_text += f"Page {i + 1}:\n{page_text}\n\n"

                    if len(ocr_text.strip()) > len(text.strip()):
                        text = ocr_text
                        logging.info("OCR provided better text extraction")

            except Exception as e:
                logging.warning(f"OCR fallback failed: {e}")

        return text
    except Exception as e:
        logging.error(f"PDF text extraction failed: {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """Extract text from DOC/DOCX files."""
    try:
        if file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text

            return text
        else:
            # For .doc files, you might need python-docx2txt or other libraries
            logging.warning("DOC format support limited. Consider converting to DOCX.")
            return ""

    except Exception as e:
        logging.error(f"DOC text extraction failed: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text files."""
    try:
        encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # If all encodings fail, try with error handling
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    except Exception as e:
        logging.error(f"Text file reading failed: {e}")
        return ""


# -------------------------------
# 4. INTELLIGENT GEMINI CLASSIFIER
# -------------------------------
def classify_document_with_gemini(text: str, filename: str = "") -> Dict[str, any]:
    """Enhanced Gemini classifier that includes document validation for specific document types."""

    model = genai.GenerativeModel('gemini-1.5-flash')

    prompt = f"""
You are an expert document analysis AI with comprehensive knowledge of global document validation standards. Analyze the provided text and intelligently determine what type of document this is, along with thorough validation analysis.

ANALYSIS INSTRUCTIONS:
1. Read through the text carefully and identify key patterns, structures, and content
2. Determine the most specific and accurate document type based on:
   - Content structure and formatting
   - Official headers, letterheads, or seals
   - Legal language and terminology
   - Reference numbers, codes, or identifiers
   - Purpose and context of the document
   - Industry-specific jargon or technical terms

3. Be as specific as possible (e.g., "Vehicle Registration Certificate" instead of just "Certificate")
4. If it's a combination document, identify the primary purpose
5. Consider cultural and regional document types (Indian, US, European, etc.)

COMPREHENSIVE VALIDATION INSTRUCTIONS:
For ALL identified documents, perform validation analysis based on document type:

IDENTITY DOCUMENTS (Aadhaar, PAN, Passport, Driver's License, Voter ID, etc.):
- Verify identification numbers follow correct format patterns and checksums
- Check date formats and logical consistency (DOB, issue dates, expiry dates)
- Validate name formatting and consistency across fields
- Analyze address formatting and completeness
- Look for proper security features mentions (holograms, watermarks, etc.)

FINANCIAL DOCUMENTS (Bank Statements, Credit Cards, Insurance, etc.):
- Verify account numbers, routing numbers, policy numbers format
- Check date sequences and transaction patterns
- Validate monetary amounts formatting
- Analyze institution names and branch codes
- Look for proper financial regulatory compliance indicators

GOVERNMENT DOCUMENTS (Certificates, Licenses, Permits, etc.):
- Verify registration numbers, certificate numbers format
- Check issuing authority details and authenticity
- Validate official seals, stamps, signatures mentions
- Analyze legal language and regulatory compliance
- Check expiry dates and renewal requirements

EDUCATIONAL DOCUMENTS (Degrees, Transcripts, Certificates, etc.):
- Verify institution names and accreditation
- Check grade formats and academic standards
- Validate course codes and credit systems
- Analyze authentication seals and signatures

BUSINESS DOCUMENTS (Registration, Tax, Incorporation, etc.):
- Verify business registration numbers and formats
- Check tax identification numbers
- Validate corporate structure information
- Analyze regulatory compliance indicators

MEDICAL DOCUMENTS (Prescriptions, Reports, Certificates, etc.):
- Verify medical professional credentials
- Check prescription formats and drug codes
- Validate medical terminology usage
- Analyze institutional authenticity

SPECIFIC VALIDATION EXAMPLES:
- PAN Card: Verify ABCDE1234F format, check name-PAN consistency
- Aadhaar: Verify 12-digit format with proper checksum, validate demographic data
- Passport: Check passport number format, validity dates, issuing authority
- Driver's License: Verify license number format, vehicle class codes, validity
- Bank Statements: Check IFSC codes, account number formats, transaction patterns
- Academic Certificates: Verify university authenticity, degree formats, grading systems

CONTEXT:
- Filename: {filename}
- Document appears to be from: [Infer from content]

TEXT TO ANALYZE:
{text[:6000]}

Please provide your analysis in this JSON format:
{{
    "document_type": "Specific document type name",
    "confidence_score": 85,
    "document_category": "Broad category (e.g., Legal, Financial, Identity, Educational, Medical, Business, etc.)",
    "country_region": "Likely country/region of origin",
    "key_identifiers": ["List of key elements that led to this classification"],
    "document_purpose": "Primary purpose or use of this document",
    "issuing_authority": "Who likely issued this document (if applicable)",
    "alternative_types": ["Other possible document types if confidence < 90"],
    "validity_status": "Valid/Invalid/Suspicious/Cannot_Determine",
    "validity_score": 85,
    "validity_issues": ["List any validation concerns, format errors, or red flags found"],
    "authenticity_indicators": ["Positive indicators of document authenticity and proper formatting"],
    "data_quality_assessment": "Overall assessment of data quality, consistency, and completeness",
    "critical_fields_analysis": {{
        "identification_numbers": "Analysis of ID numbers, account numbers, registration numbers",
        "dates_and_validity": "Analysis of dates, expiry, consistency",
        "personal_information": "Analysis of names, addresses, demographic data",
        "institutional_details": "Analysis of issuing authority, signatures, seals"
    }},
    "compliance_indicators": ["Regulatory compliance and standard format adherence"],
    "security_features": ["Mentioned security features like holograms, watermarks, etc."],
    "recommendation": "Overall recommendation for document acceptance/rejection"
}}

IMPORTANT: 
- Return only the JSON response, no additional text
- For validity_status: use "Valid" if document data appears authentic, "Invalid" if clearly fake/wrong, "Suspicious" if questionable, "Cannot_Determine" if insufficient info
- Be thorough in your validation analysis, especially for government documents
"""

    try:
        response = model.generate_content(prompt)
        result_text = response.text.strip()

        # Clean up the response to extract JSON
        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            result = json.loads(json_str)
        else:
            # Fallback parsing if JSON format is not perfect
            result = parse_enhanced_text_response(result_text)

        # Ensure all required fields are present
        required_fields = [
            "validity_status", "validity_score", "validity_issues",
            "authenticity_indicators", "data_quality_assessment",
            "critical_fields_analysis", "compliance_indicators",
            "security_features", "recommendation"
        ]

        for field in required_fields:
            if field not in result:
                result[field] = get_default_value(field)

        return result

    except json.JSONDecodeError as e:
        logging.warning(f"JSON parsing failed, attempting text parsing: {e}")
        return parse_enhanced_text_response(response.text.strip())

    except Exception as e:
        logging.error(f"Gemini classification failed: {e}")
        return get_error_result(str(e))


def parse_enhanced_text_response(response_text: str) -> Dict[str, any]:
    """Enhanced fallback parser for non-JSON responses that includes validation fields."""
    result = get_default_result()

    try:
        lines = response_text.split('\n')

        for line in lines:
            line = line.strip()
            if any(keyword in line.lower() for keyword in ['document type', 'type:']):
                result["document_type"] = extract_value_after_colon(line)
            elif 'confidence' in line.lower():
                conf_match = re.search(r'\d+', line)
                if conf_match:
                    result["confidence_score"] = int(conf_match.group())
            elif 'category' in line.lower():
                result["document_category"] = extract_value_after_colon(line)
            elif 'country' in line.lower() or 'region' in line.lower():
                result["country_region"] = extract_value_after_colon(line)
            elif 'purpose' in line.lower():
                result["document_purpose"] = extract_value_after_colon(line)
            elif 'authority' in line.lower() or 'issuer' in line.lower():
                result["issuing_authority"] = extract_value_after_colon(line)
            elif 'validity' in line.lower() and 'status' in line.lower():
                result["validity_status"] = extract_value_after_colon(line)
            elif 'validity' in line.lower() and 'score' in line.lower():
                score_match = re.search(r'\d+', line)
                if score_match:
                    result["validity_score"] = int(score_match.group())

    except Exception as e:
        logging.warning(f"Enhanced text parsing failed: {e}")
        result["validity_issues"].append(f"Parsing error: {str(e)}")

    return result


def extract_value_after_colon(text: str) -> str:
    """Extract value after colon in a line."""
    if ':' in text:
        return text.split(':', 1)[1].strip().strip('"\'')
    return text.strip()


def get_default_result() -> Dict[str, any]:
    """Get default result structure."""
    return {
        "document_type": "Unknown Document",
        "confidence_score": 0,
        "document_category": "Unknown",
        "country_region": "Unknown",
        "key_identifiers": [],
        "document_purpose": "Unknown",
        "issuing_authority": "Unknown",
        "alternative_types": [],
        "validity_status": "Cannot_Determine",
        "validity_score": 0,
        "validity_issues": [],
        "authenticity_indicators": [],
        "data_quality_assessment": "Could not parse response properly",
        "critical_fields_analysis": {
            "identification_numbers": "Not analyzed",
            "dates_and_validity": "Not analyzed",
            "personal_information": "Not analyzed",
            "institutional_details": "Not analyzed"
        },
        "compliance_indicators": [],
        "security_features": [],
        "recommendation": "Manual review recommended"
    }


def get_error_result(error_msg: str) -> Dict[str, any]:
    """Get error result structure."""
    result = get_default_result()
    result["validity_issues"] = [f"Classification error: {error_msg}"]
    result["data_quality_assessment"] = f"Error occurred: {error_msg}"
    result["recommendation"] = "Manual review required due to processing error"
    return result


def get_default_value(field: str):
    """Get default values for missing fields."""
    defaults = {
        "validity_status": "Cannot_Determine",
        "validity_score": 0,
        "validity_issues": [],
        "authenticity_indicators": [],
        "data_quality_assessment": "No assessment available",
        "critical_fields_analysis": {
            "identification_numbers": "Not analyzed",
            "dates_and_validity": "Not analyzed",
            "personal_information": "Not analyzed",
            "institutional_details": "Not analyzed"
        },
        "compliance_indicators": [],
        "security_features": [],
        "recommendation": "Manual review recommended"
    }
    return defaults.get(field, "Unknown")


# -------------------------------
# 5. DOCUMENT INSIGHTS
# -------------------------------
def get_document_insights(text: str, classification: Dict) -> Dict[str, any]:
    """Get additional insights about the document using Gemini."""
    model = genai.GenerativeModel('gemini-1.5-flash')

    validity_context = f"""
    Document Validity Status: {classification.get('validity_status', 'Unknown')}
    Validity Score: {classification.get('validity_score', 0)}%
    Identified Issues: {', '.join(classification.get('validity_issues', []))}
    Authenticity Indicators: {', '.join(classification.get('authenticity_indicators', []))}
    """

    prompt = f"""
Based on the document classification and validation analysis, provide additional insights:

DOCUMENT TYPE: {classification['document_type']}
VALIDATION CONTEXT: {validity_context}
DOCUMENT TEXT: {text[:3000]}

Please analyze and provide:
1. Detailed explanation of the validation findings
2. Potential reasons for any identified issues
3. Recommendations for document verification
4. Legal or compliance implications (if any)
5. Next steps or actions recommended
6. How to verify authenticity independently
7.tell whether to accept or reject the document based on validity result


Focus especially on validation concerns and provide actionable advice.
Provide insights in a clear, structured format.
"""

    try:
        response = model.generate_content(prompt)
        return {
            "detailed_insights": response.text.strip(),
            "validation_summary": {
                "status": classification.get('validity_status', 'Unknown'),
                "score": classification.get('validity_score', 0),
                "main_concerns": classification.get('validity_issues', []),
                "positive_indicators": classification.get('authenticity_indicators', [])
            }
        }
    except Exception as e:
        return {"detailed_insights": f"Unable to generate insights: {str(e)}"}


# -------------------------------
# 6. MAIN ANALYSIS FUNCTION
# -------------------------------
def analyze_document(file_path: str, api_key: str = None, get_insights: bool = True) -> Dict[str, any]:
    """Main function to analyze a document file."""

    try:
        # Configure Gemini API
        configure_gemini(api_key)

        # Extract text from file
        logging.info(f"Extracting text from: {file_path}")
        text = extract_text_from_file(file_path)

        if not text.strip():
            return {
                "error": "No text could be extracted from the document",
                "file_path": file_path
            }

        # Classify and validate document
        logging.info("Analyzing document with Gemini...")
        classification = classify_document_with_gemini(text, os.path.basename(file_path))

        # Get additional insights if requested
        insights = None
        if get_insights:
            logging.info("Generating additional insights...")
            insights = get_document_insights(text, classification)

        # Prepare final result
        result = {
            "file_path": file_path,
            "file_type": Path(file_path).suffix.lower(),
            "text_length": len(text),
            "extracted_text_preview": text[:500] + "..." if len(text) > 500 else text,
            "classification": classification,
            "insights": insights,
            "processing_status": "success"
        }

        return result

    except Exception as e:
        logging.error(f"Document analysis failed: {e}")
        return {
            "error": str(e),
            "file_path": file_path,
            "processing_status": "failed"
        }


# -------------------------------
# 7. BATCH PROCESSING
# -------------------------------
def analyze_multiple_documents(input_path: str, api_key: str = None,
                               get_insights: bool = True, output_dir: str = None) -> List[Dict[str, any]]:
    """Analyze multiple documents from a directory or pattern."""

    try:
        # Get list of files to process
        files = get_supported_files(input_path)
        print(f"Found {len(files)} files to process")

        results = []

        for i, file_path in enumerate(files, 1):
            print(f"\n📄 Processing file {i}/{len(files)}: {os.path.basename(file_path)}")

            try:
                result = analyze_document(file_path, api_key, get_insights)
                results.append(result)

                # Save individual result if output directory provided
                if output_dir:
                    os.makedirs(output_dir, exist_ok=True)
                    filename = Path(file_path).stem + "_analysis.json"
                    output_file = os.path.join(output_dir, filename)

                    with open(output_file, 'w', encoding='utf-8') as f:
                        json.dump(result, f, indent=2, ensure_ascii=False)

                    print(f"💾 Individual result saved to: {output_file}")

            except Exception as e:
                error_result = {
                    "error": str(e),
                    "file_path": file_path,
                    "processing_status": "failed"
                }
                results.append(error_result)
                print(f"❌ Failed to process {file_path}: {e}")

        return results

    except Exception as e:
        print(f"❌ Batch processing failed: {e}")
        return []


# -------------------------------
# 8. PRETTY PRINTING FUNCTIONS
# -------------------------------
def print_analysis_results(result: Dict[str, any]):
    """Print analysis results in a formatted way."""

    print("=" * 80)
    print("📄 DOCUMENT ANALYSIS RESULTS")
    print("=" * 80)

    if "error" in result:
        print(f"❌ Error: {result['error']}")
        return

    # Basic file info
    print(f"📁 File: {result['file_path']}")
    print(f"📋 File Type: {result['file_type']}")
    print(f"📝 Text Length: {result['text_length']} characters")

    # Classification results
    classification = result['f']
    print(f"\n🎯 DOCUMENT CLASSIFICATION")
    print("-" * 40)
    print(f"Document Type: {classification['document_type']}")
    print(f"Category: {classification['document_category']}")
    print(f"Country/Region: {classification['country_region']}")
    print(f"Confidence Score: {classification['confidence_score']}%")
    print(f"Issuing Authority: {classification['issuing_authority']}")

    # Validity analysis
    print(f"\n✅ VALIDITY ANALYSIS")
    print("-" * 40)
    print(f"Validity Status: {classification['validity_status']}")
    print(f"Validity Score: {classification['validity_score']}%")
    print(f"Recommendation: {classification['recommendation']}")

    if classification.get('validity_issues'):
        print(f"\n⚠️  ISSUES IDENTIFIED:")
        for issue in classification['validity_issues'][:3]:
            print(f"  • {issue}")

    if classification.get('authenticity_indicators'):
        print(f"\n✨ POSITIVE INDICATORS:")
        for indicator in classification['authenticity_indicators'][:3]:
            print(f"  • {indicator}")

    # Critical fields analysis
    critical_analysis = classification.get('critical_fields_analysis', {})
    if any(critical_analysis.values()):
        print(f"\n🔍 CRITICAL FIELDS ANALYSIS")
        print("-" * 40)
        for field, analysis in critical_analysis.items():
            if analysis and analysis != "Not analyzed":
                print(f"{field.replace('_', ' ').title()}: {analysis}")

    # Text preview
    print(f"\n📖 TEXT PREVIEW")
    print("-" * 40)
    print(result['extracted_text_preview'])

    # Insights
    if result.get('insights'):
        print(f"\n💡 ADDITIONAL INSIGHTS")
        print("-" * 40)
        print(result['insights']['detailed_insights'])


def print_batch_summary(results: List[Dict[str, any]]):
    """Print summary of batch processing results."""

    print("\n" + "=" * 80)
    print("📊 BATCH PROCESSING SUMMARY")
    print("=" * 80)

    total_files = len(results)
    successful = len([r for r in results if r.get('processing_status') == 'success'])
    failed = total_files - successful

    print(f"Total Files Processed: {total_files}")
    print(f"Successfully Analyzed: {successful}")
    print(f"Failed: {failed}")

    if successful > 0:
        print(f"\n📈 DOCUMENT TYPE DISTRIBUTION:")
        doc_types = {}
        validity_stats = {"Valid": 0, "Invalid": 0, "Suspicious": 0, "Cannot_Determine": 0}

        for result in results:
            if result.get('processing_status') == 'success':
                doc_type = result['classification']['document_type']
                doc_types[doc_type] = doc_types.get(doc_type, 0) + 1

                validity = result['classification']['validity_status']
                validity_stats[validity] = validity_stats.get(validity, 0) + 1

        for doc_type, count in sorted(doc_types.items()):
            print(f"  • {doc_type}: {count}")

        print(f"\n🔍 VALIDITY DISTRIBUTION:")
        for status, count in validity_stats.items():
            if count > 0:
                print(f"  • {status}: {count}")


# -------------------------------
# 9. COMMAND LINE INTERFACE
# -------------------------------
def main():
    """Document analysis with direct input configuration."""

    # Configuration - Replace command line arguments with direct input
    config = {
        'input_path': "D:\\demo\\sample_aadhaar.pdf",  # Path to document file, directory, or glob pattern
        'api_key': "AIzaSyABg9-HiG1wndO8daIZFVKC9EPdwCTe3Is",  # Gemini API key (or set GEMINI_API_KEY environment variable)
        'no_insights': False,  # Skip generating additional insights
        'output': None,  # Output file path for JSON results (single file) or directory (batch)
        'batch': False,  # Process multiple files (auto-detected for directories)
        'quiet': False  # Suppress detailed output, show only summary
    }

    # You can modify the config above or uncomment and modify these individual variables:
    # input_path = 'documents/sample.pdf'
    # api_key = 'your_gemini_api_key_here'
    # no_insights = False
    # output_path = 'results.json'
    # batch_mode = False
    # quiet_mode = False

    # Use config values
    input_path = config['input_path']
    api_key = config['api_key']
    no_insights = config['no_insights']
    output = config['output']
    batch = config['batch']
    quiet = config['quiet']

    try:
        # Determine if this is batch processing
        is_batch = batch or os.path.isdir(input_path) or '*' in input_path

        if is_batch:
            # Batch processing
            print(f"🔄 Starting batch processing for: {input_path}")
            results = analyze_multiple_documents(
                input_path=input_path,
                api_key=api_key,
                get_insights=not no_insights,
                output_dir=output if output else None
            )

            if not quiet:
                for result in results:
                    print_analysis_results(result)
                    print("\n" + "-" * 80 + "\n")

            print_batch_summary(results)

            # Save combined results if output specified and not a directory
            if output and not os.path.isdir(output):
                with open(output, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
                print(f"\n💾 Combined results saved to: {output}")

        else:
            # Single file processing
            if not os.path.exists(input_path):
                print(f"❌ Error: File not found: {input_path}")
                return False

            print(f"🔄 Analyzing: {input_path}")
            result = analyze_document(
                file_path=input_path,
                api_key=api_key,
                get_insights=not no_insights
            )

            # Print results
            if not quiet:
                print_analysis_results(result)

            # Save to file if requested
            if output:
                with open(output, 'w', encoding='utf-8') as f:
                    json.dump(result, f, indent=2, ensure_ascii=False)
                print(f"\n💾 Results saved to: {output}")

        return True

    except Exception as e:
        print(f"❌ Error: {str(e)}")
        return False


# Example usage configurations:

def analyze_single_document_example():
    """Example: Analyze a single document with insights."""
    config = {
        'input_path': 'contract.pdf',
        'api_key': None,  # Uses GEMINI_API_KEY environment variable
        'no_insights': False,
        'output': 'contract_analysis.json',
        'batch': False,
        'quiet': False
    }

    # Update the main config and run
    globals()['config'] = config
    main()


def batch_process_example():
    """Example: Batch process all PDFs in a directory."""
    config = {
        'input_path': 'documents/',
        'api_key': 'your_api_key_here',
        'no_insights': True,  # Skip insights for faster processing
        'output': 'batch_results/',
        'batch': True,
        'quiet': True
    }

    globals()['config'] = config
    main()


def quick_analysis_example():
    """Example: Quick analysis without saving results."""
    config = {
        'input_path': 'important_doc.pdf',
        'api_key': None,
        'no_insights': True,
        'output': None,  # Don't save, just print
        'batch': False,
        'quiet': False
    }

    globals()['config'] = config
    main()


# Run the main function
if __name__ == "__main__":
    main()